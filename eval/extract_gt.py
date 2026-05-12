#!/usr/bin/env python3
"""Extract ground-truth images and captions from Rov_Immagine.docx.

The docx contains 93 captioned figures across 11 ROV videos. This script:
  1. Pulls every embedded image out of word/media/ via zipfile.
  2. Saves them to eval/gt_images/figure_<NN>.jpg.
  3. Walks doc.paragraphs to pair each image with the next non-empty
     caption paragraph and parses the caption.
  4. Writes a full template (all 93) and a 20-figure subset for
     hand-annotation.

The docx itself is private survey data (kept outside the repo); the
extracted images are gitignored. The two JSON files ARE tracked — they
are reproducible from the docx and form the eval contract.

Usage:
    python eval/extract_gt.py --docx ../Rov_Immagine.docx
"""
from __future__ import annotations
import argparse
import json
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document


CAPTION_RE = re.compile(
    r"Figura\s*(\d+)\s*[-–]\s*Messina\s*[-–]?\s*video\s*(\d+)\s*"
    r"Foto\s*(\d+)\s*[-–]\s*(.+)",
    re.IGNORECASE,
)

SUBSET_IDS = [1, 3, 5, 7, 9, 11, 13, 14, 17, 18, 20, 21, 24, 25, 27, 33, 38, 50, 70, 85]

EMPTY_EXPECTED = {
    "tipo_fondale": None,
    "granulometria": None,
    "presenza_rocce": None,
    "presenza_ciottoli": None,
    "copertura_algale": None,
    "taxa_algali": [],
    "elementi_antropici": [],
    "fauna": [],
}


def extract_images(docx_path: Path, out_dir: Path) -> list[Path]:
    """Pull every image from word/media/ inside the docx, sorted by name.

    Returns the list of saved paths in the order the images appear in the
    archive (which matches in-document order for Word's default layout).
    """
    out_dir.mkdir(parents=True, exist_ok=True)
    for old in out_dir.glob("figure_*.jpg"):
        old.unlink()

    saved: list[Path] = []
    with zipfile.ZipFile(docx_path) as z:
        media = sorted(
            n for n in z.namelist()
            if n.startswith("word/media/") and not n.endswith("/")
        )
        for idx, name in enumerate(media, start=1):
            dst = out_dir / f"figure_{idx:02d}.jpg"
            with z.open(name) as src, dst.open("wb") as out:
                shutil.copyfileobj(src, out)
            saved.append(dst)
    return saved


def extract_captions(docx_path: Path) -> list[str]:
    """Walk paragraphs, pair each image-bearing paragraph with the next
    non-empty caption paragraph. Returns captions in document order.
    """
    doc = Document(str(docx_path))
    captions: list[str] = []
    pending_image = False
    for para in doc.paragraphs:
        has_image = any(
            "graphic" in run._element.xml.lower() or "<w:drawing" in run._element.xml
            for run in para.runs
        )
        text = para.text.strip()
        if has_image:
            pending_image = True
            continue
        if pending_image and text:
            captions.append(text)
            pending_image = False
    return captions


def parse_caption(caption: str) -> tuple[int | None, int | None, int | None, str]:
    m = CAPTION_RE.search(caption)
    if not m:
        return None, None, None, caption.strip()
    fig = int(m.group(1))
    video = int(m.group(2))
    photo = int(m.group(3))
    desc = m.group(4).strip()
    return fig, video, photo, desc


def build_entries(image_paths: list[Path], captions: list[str]) -> list[dict]:
    entries: list[dict] = []
    n = min(len(image_paths), len(captions))
    if len(image_paths) != len(captions):
        print(
            f"WARN: {len(image_paths)} images vs {len(captions)} captions — "
            f"pairing the first {n}."
        )
    for i in range(n):
        img_path = image_paths[i]
        caption = captions[i]
        fig_id, video, photo, desc = parse_caption(caption)
        if fig_id is None:
            fig_id = i + 1
            print(f"WARN: could not parse caption for figure {fig_id}: {caption!r}")
        entry = {
            "figure_id": fig_id,
            "video": video,
            "photo": photo,
            "image_path": f"eval/gt_images/{img_path.name}",
            "original_caption": caption,
            "expected": dict(EMPTY_EXPECTED, taxa_algali=[], elementi_antropici=[], fauna=[]),
            "annotator_notes": "",
        }
        entries.append(entry)
    return entries


def write_json(entries: list[dict], path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(entries, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def main() -> None:
    p = argparse.ArgumentParser(description=__doc__.splitlines()[0])
    p.add_argument(
        "--docx",
        type=Path,
        default=Path("../Rov_Immagine.docx"),
        help="Path to Rov_Immagine.docx (default: ../Rov_Immagine.docx).",
    )
    p.add_argument(
        "--eval-dir",
        type=Path,
        default=Path("eval"),
        help="Eval directory root (default: eval).",
    )
    args = p.parse_args()

    if not args.docx.exists():
        raise FileNotFoundError(
            f"docx not found at {args.docx}. Pass --docx to point at it."
        )

    images_dir = args.eval_dir / "gt_images"
    image_paths = extract_images(args.docx, images_dir)
    print(f"Extracted {len(image_paths)} images to {images_dir}")

    captions = extract_captions(args.docx)
    print(f"Parsed {len(captions)} captions from the docx")

    entries = build_entries(image_paths, captions)
    write_json(entries, args.eval_dir / "gt_annotations_template.json")
    print(f"Wrote {len(entries)} entries to gt_annotations_template.json")

    subset = [e for e in entries if e["figure_id"] in SUBSET_IDS]
    write_json(subset, args.eval_dir / "gt_annotations_subset.json")
    print(
        f"Wrote {len(subset)} entries to gt_annotations_subset.json "
        f"(target: {len(SUBSET_IDS)})"
    )


if __name__ == "__main__":
    main()
